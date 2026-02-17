"""
Multimedia Lecture Processor

Processes multimedia lecture materials (PDF slide decks, VTT video transcripts,
and slide timing data) into accessible Word documents. The output combines slide
images with their corresponding transcript text.

Input files (all in multimedia_materials/):
- One PDF file (slide deck)
- One VTT file (video transcript)
- One TXT file (slide timestamps in MM:SS format)

Output: Word document in multimedia_finished_transcripts/ with MULTIMEDIA_ prefix
"""

import os
from io import BytesIO
from pathlib import Path

from pdf2image import convert_from_path

from PIL import Image
from docx import Document
from docx.shared import Inches

INPUT_DIR = "multimedia_materials"
OUTPUT_DIR = "multimedia_finished_transcripts"
OUTPUT_PREFIX = "MULTIMEDIA_"


def discover_input_files(input_dir: str) -> tuple[str, str, str]:
    """
    Find exactly 1 PDF, 1 VTT, 1 TXT file in the input directory.

    Args:
        input_dir: Directory to search for input files

    Returns:
        Tuple of (pdf_path, vtt_path, txt_path)

    Raises:
        ValueError: If validation fails (wrong number of files)
    """
    if not os.path.exists(input_dir):
        raise ValueError(f"Error: Input directory '{input_dir}/' does not exist")

    files = os.listdir(input_dir)

    # Find files by extension
    pdf_files = [f for f in files if f.lower().endswith('.pdf')]
    vtt_files = [f for f in files if f.lower().endswith('.vtt')]
    txt_files = [f for f in files if f.lower().endswith('.txt')]

    # Validate counts
    if len(pdf_files) == 0:
        raise ValueError(f"Error: No PDF file found in {input_dir}/")
    if len(pdf_files) > 1:
        raise ValueError(
            f"Error: Multiple PDF files found: {pdf_files}. Expected exactly 1."
        )

    if len(vtt_files) == 0:
        raise ValueError(f"Error: No VTT file found in {input_dir}/")
    if len(vtt_files) > 1:
        raise ValueError(
            f"Error: Multiple VTT files found: {vtt_files}. Expected exactly 1."
        )

    if len(txt_files) == 0:
        raise ValueError(f"Error: No TXT file found in {input_dir}/")
    if len(txt_files) > 1:
        raise ValueError(
            f"Error: Multiple TXT files found: {txt_files}. Expected exactly 1."
        )

    return (
        os.path.join(input_dir, pdf_files[0]),
        os.path.join(input_dir, vtt_files[0]),
        os.path.join(input_dir, txt_files[0])
    )


def extract_pdf_slides(pdf_path: str) -> list[Image.Image]:
    """
    Extract slides from PDF as images.

    Args:
        pdf_path: Path to PDF file

    Returns:
        List of PIL Images, one per slide

    Raises:
        ValueError: If PDF is empty or invalid
    """
    try:
        slides = convert_from_path(pdf_path, dpi=150, fmt='PNG')
    except Exception as e:
        if "poppler" in str(e).lower():
            raise ImportError(
                "Error: poppler not found. Install with:\n"
                "  Windows: Download from https://github.com/oschwartz10612/poppler-windows/releases/\n"
                "           Extract and add the 'bin' folder to your PATH\n"
                "           OR: conda install -c conda-forge poppler (if using conda)\n"
                "  macOS: brew install poppler\n"
                "  Linux: sudo apt-get install poppler-utils"
            ) from e
        raise ValueError(f"Error: Failed to extract slides from PDF: {e}") from e

    if not slides:
        raise ValueError(f"Error: PDF file '{pdf_path}' contains no pages")

    return slides


def vtt_timestamp_to_seconds(timestamp: str) -> float:
    """
    Convert VTT timestamp (HH:MM:SS.mmm) to seconds.

    Args:
        timestamp: VTT timestamp string

    Returns:
        Time in seconds (float)
    """
    time_part, ms_part = timestamp.strip().split('.')
    h, m, s = map(int, time_part.split(':'))
    ms = int(ms_part)

    return h * 3600 + m * 60 + s + ms / 1000.0


def parse_vtt_with_timestamps(vtt_path: str) -> list[dict]:
    """
    Parse VTT file and preserve timestamps.

    VTT structure:
    - First 2 lines: WEBVTT header
    - Chunks separated by blank lines (\n\n)
    - Each chunk: entry number, timestamp range, text (with possible <v -> tags)

    Args:
        vtt_path: Path to VTT file

    Returns:
        List of dicts: [{"start": 3.39, "end": 7.17, "text": "..."}]
    """
    with open(vtt_path, 'r', encoding='utf-8') as f:
        # Skip first 2 lines (WEBVTT header)
        next(f)
        next(f)

        # Read remaining content
        file_content = f.read()

    # Split on blank lines
    chunks = file_content.split('\n\n')

    entries = []
    for chunk in chunks:
        if not chunk.strip():
            continue

        lines = chunk.split('\n')
        if len(lines) < 3:
            continue  # Skip malformed chunks

        # Line 1: timestamp range (HH:MM:SS.mmm --> HH:MM:SS.mmm)
        timestamp_line = lines[1]
        if '-->' not in timestamp_line:
            continue  # Skip if no timestamp

        start_str, end_str = timestamp_line.split('-->')
        start_time = vtt_timestamp_to_seconds(start_str.strip())
        end_time = vtt_timestamp_to_seconds(end_str.strip())

        # Line 2+: text (remove <v -> tags)
        text = ' '.join(lines[2:])
        text = text.replace('<v ->', '').strip()

        entries.append({
            'start': start_time,
            'end': end_time,
            'text': text
        })

    return entries


def mmss_to_seconds(timestamp: str) -> float:
    """
    Convert timestamp to seconds. Supports HH:MM:SS or MM:SS format.

    Args:
        timestamp: Timestamp in HH:MM:SS, MM:SS, H:MM:SS, or M:SS format
                   If HH is omitted, defaults to 0 hours

    Returns:
        Time in seconds (float)

    Raises:
        ValueError: If format is invalid
    """
    parts = timestamp.strip().split(':')

    if len(parts) == 2:
        # MM:SS format - treat as 0:MM:SS
        hours = 0
        minutes_str, seconds_str = parts
    elif len(parts) == 3:
        # HH:MM:SS format
        hours_str, minutes_str, seconds_str = parts
        try:
            hours = int(hours_str)
        except ValueError:
            raise ValueError(f"Invalid format: '{timestamp}'. Expected numeric values")

        if hours < 0:
            raise ValueError(f"Invalid values: '{timestamp}'. Hours must be >= 0")
    else:
        raise ValueError(
            f"Invalid format: '{timestamp}'. Expected HH:MM:SS or MM:SS"
        )

    try:
        minutes = int(minutes_str)
        seconds = int(seconds_str)
    except ValueError:
        raise ValueError(f"Invalid format: '{timestamp}'. Expected numeric values")

    if seconds >= 60 or seconds < 0 or minutes < 0:
        raise ValueError(
            f"Invalid values: '{timestamp}'. Minutes and seconds must be >= 0, "
            f"seconds must be < 60"
        )

    if minutes >= 60:
        raise ValueError(
            f"Invalid values: '{timestamp}'. Minutes must be < 60"
        )

    return hours * 3600 + minutes * 60 + seconds


def parse_timestamp_file(txt_path: str) -> list[float]:
    """
    Parse timestamp file with HH:MM:SS or MM:SS format.

    Validates:
    - Format: HH:MM:SS or MM:SS (e.g., "01:30" -> 90.0, "1:05:30" -> 3930.0)
    - If HH is omitted, defaults to 0 hours
    - Timestamps in ascending order
    - First timestamp must be 00:00 or 0:00:00

    Args:
        txt_path: Path to timestamp file

    Returns:
        List of timestamps in seconds

    Raises:
        ValueError: If validation fails
    """
    timestamps = []

    with open(txt_path, 'r', encoding='utf-8') as f:
        for line_num, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue  # Skip empty lines

            try:
                timestamp = mmss_to_seconds(line)
            except ValueError as e:
                raise ValueError(f"Error at line {line_num}: {e}") from e

            timestamps.append(timestamp)

    if not timestamps:
        raise ValueError("Error: Timestamp file is empty")

    # Validate first timestamp is 00:00 or 0:00:00
    if timestamps[0] != 0.0:
        raise ValueError(
            f"Error: First timestamp must be 00:00 or 0:00:00, got {timestamps[0]:.1f} seconds"
        )

    # Validate ascending order
    for i in range(1, len(timestamps)):
        if timestamps[i] <= timestamps[i - 1]:
            raise ValueError(
                f"Error: Timestamps must be in ascending order. "
                f"Timestamp at line {i + 1} ({timestamps[i]:.1f}s) is not greater than "
                f"previous ({timestamps[i - 1]:.1f}s)"
            )

    return timestamps


def validate_alignment(
    slide_count: int,
    timestamps: list[float],
    vtt_entries: list[dict]
) -> None:
    """
    Validate alignment between slides, timestamps, and VTT entries.

    Validates:
    - Slide count matches timestamp count
    - All timestamps <= max VTT end time
    - Timestamps strictly ascending (already validated in parse)
    - First timestamp == 0.0 (already validated in parse)

    Args:
        slide_count: Number of slides in PDF
        timestamps: List of slide timestamps
        vtt_entries: List of VTT entries with start/end times

    Raises:
        ValueError: If validation fails
    """
    # Validate count match
    if len(timestamps) != slide_count:
        raise ValueError(
            f"Error: Slide count ({slide_count}) does not match "
            f"timestamp count ({len(timestamps)})"
        )

    # Get max video duration
    if not vtt_entries:
        raise ValueError("Error: VTT file contains no valid entries")

    max_video_time = max(entry['end'] for entry in vtt_entries)

    # Validate all timestamps within video duration
    for i, ts in enumerate(timestamps):
        if ts > max_video_time:
            minutes = int(ts // 60)
            seconds = int(ts % 60)
            max_minutes = int(max_video_time // 60)
            max_seconds = int(max_video_time % 60)
            raise ValueError(
                f"Error: Timestamp {i + 1} ({minutes:02d}:{seconds:02d}) exceeds "
                f"video duration ({max_minutes:02d}:{max_seconds:02d})"
            )


def reconstruct_sentences(text_chunks: list[str]) -> str:
    """
    Reconstruct sentences from text chunks.

    Concatenates text chunks until terminal punctuation (. ! ?) is found.
    Adapted from vtt_to_transcript.py.

    Args:
        text_chunks: List of text fragments

    Returns:
        String with properly reconstructed sentences
    """
    sentences = []
    current_sentence = ""
    sentence_punctuations = ['.', '!', '?']

    for text in text_chunks:
        # Skip empty text
        if not text:
            continue

        # If ending punctuation is found, add to current sentence and finalize
        if text[-1] in sentence_punctuations:
            current_sentence += text + " "
            sentences.append(current_sentence)
            current_sentence = ""
        else:
            # Add the text with a space for the next text
            current_sentence += text + " "

    # Add any remaining text
    if current_sentence:
        sentences.append(current_sentence)

    return ''.join(sentences).strip()


def segment_transcript(
    vtt_entries: list[dict],
    timestamps: list[float]
) -> list[str]:
    """
    Segment transcript by slide timestamps.

    For each slide timestamp, collect VTT entries in range:
    - Slide i: timestamps[i] <= entry['start'] < timestamps[i+1]
    - Last slide: timestamps[-1] <= entry['start'] < infinity

    Args:
        vtt_entries: List of VTT entries with start/end times and text
        timestamps: List of slide timestamps in seconds

    Returns:
        List of transcript segments (one per slide)
    """
    segments = []

    for i in range(len(timestamps)):
        # Determine time range for this slide
        start_time = timestamps[i]
        end_time = timestamps[i + 1] if i + 1 < len(timestamps) else float('inf')

        # Collect VTT entries in this range
        text_chunks = []
        for entry in vtt_entries:
            if start_time <= entry['start'] < end_time:
                text_chunks.append(entry['text'])

        # Reconstruct sentences
        if text_chunks:
            segment_text = reconstruct_sentences(text_chunks)
            segments.append(segment_text)
        else:
            segments.append("")  # Empty segment

    return segments


def create_word_document(
    slides: list[Image.Image],
    transcript_segments: list[str],
    output_path: str
) -> None:
    """
    Create Word document with alternating slide images and transcript segments.

    Format:
    - Image (6 inches wide, aspect ratio preserved)
    - Paragraph spacing
    - Transcript text or "(No transcript for this section)" in italics
    - Page break (except after last slide)

    Args:
        slides: List of PIL Images (slide images)
        transcript_segments: List of transcript text (one per slide)
        output_path: Path to save Word document
    """
    doc = Document()

    for i, (slide_image, transcript_text) in enumerate(zip(slides, transcript_segments)):
        # Add slide image
        # Convert PIL Image to BytesIO stream
        img_stream = BytesIO()
        slide_image.save(img_stream, format='PNG')
        img_stream.seek(0)

        doc.add_picture(img_stream, width=Inches(6))

        # Add spacing
        doc.add_paragraph()

        # Add transcript text
        if transcript_text:
            doc.add_paragraph(transcript_text)
        else:
            # No transcript for this section
            p = doc.add_paragraph("(No transcript for this section)")
            p.runs[0].italic = True

        # Add page break (except after last slide)
        if i < len(slides) - 1:
            doc.add_page_break()

    # Save document
    doc.save(output_path)


def main():
    """Main execution function."""
    try:
        # 1. Discover and validate files
        print(f"Discovering input files in '{INPUT_DIR}/'...")
        pdf_path, vtt_path, txt_path = discover_input_files(INPUT_DIR)

        pdf_name = os.path.basename(pdf_path)
        vtt_name = os.path.basename(vtt_path)
        txt_name = os.path.basename(txt_path)

        print(f"  Found PDF: {pdf_name}")
        print(f"  Found VTT: {vtt_name}")
        print(f"  Found TXT: {txt_name}")

        # 2. Extract data
        print("\nExtracting PDF slides...")
        slides = extract_pdf_slides(pdf_path)
        print(f"  Extracted {len(slides)} slide(s)")

        print("\nParsing VTT transcript...")
        vtt_entries = parse_vtt_with_timestamps(vtt_path)
        print(f"  Parsed {len(vtt_entries)} transcript entry/entries")

        print("\nParsing timestamp file...")
        timestamps = parse_timestamp_file(txt_path)
        print(f"  Parsed {len(timestamps)} timestamp(s)")

        # 3. Validate alignment
        print("\nValidating alignment...")
        validate_alignment(len(slides), timestamps, vtt_entries)
        print("  Alignment validated successfully")

        # 4. Segment transcript by timestamps
        print("\nSegmenting transcript by slide timestamps...")
        transcript_segments = segment_transcript(vtt_entries, timestamps)
        print(f"  Created {len(transcript_segments)} transcript segment(s)")

        # 5. Create Word document
        print("\nCreating Word document...")

        # Ensure output directory exists
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Generate output filename
        pdf_basename = Path(pdf_path).stem
        output_filename = f"{OUTPUT_PREFIX}{pdf_basename}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        create_word_document(slides, transcript_segments, output_path)

        print(f"\nSuccessfully created: {output_path}")

    except (ValueError, ImportError) as e:
        print(f"\n{e}")
        return 1
    except Exception as e:
        print(f"\nUnexpected error: {e}")
        raise

    return 0


if __name__ == "__main__":
    exit(main())
